from copy import deepcopy
from math import ceil
from PIL import Image

from django.shortcuts import render
from docx.enum.table import WD_TABLE_ALIGNMENT
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Inches, Pt, Cm

import os
from docx import Document
import docx
import json

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def generate_docx(template_path, pro_info, save_path, dics:list=[]):
    tpl = DocxTemplate(template_path)

    tpl.render(pro_info)
    key = 'Name'
    key2 = 'ReportNum'
    doc_path = os.path.join(save_path, f'{pro_info[key]}评估报告-{pro_info[key2]}.docx')
    tpl.save(doc_path)

    # put doc into for
    for pip_dic in dics:
        doc = docx.Document(doc_path)
        fill_defect(document=doc, pipe_dic=pip_dic, save_path=save_path)
        doc.save(doc_path)

# context = {'Name)': '',
    #            'Address)': '',
    #            'Client)': '',
    #            'CheckPerson)': '',
    #            'StartDate)': '',
    #            'DateTime) ': '',
    #            'Author)': '',
    #            'Proofreader)': '',
    #            'Checker)': '',
    #            'Ratifier)': '',
    #            'CheckPerson)': '',
    #            'Address)': '',
    #            'SupervisionUnit)': '',
    #            'PipeCount)': '',
    #            'DistanceTotal)': '',
    #            'CheckDistance)': '',
    #            'CheckType)': '',
    #            'StartDate)': '',
    #            'EquipmentName)': '',
    #            'MoveMethon)': '',
    #            'SealedMode)': '',
    #            'DrainageMode)': '',
    #            'CleanMode)': '',
    #             }

defect_dic = {
    'distance':'/',
    'name':'',
    'mark':'/',
    'grade':'/',
    'description':'/',
    'img_path':'',
}

pipe_dic = {
    'video_name': '',
    'begin_well': '',
    'end_well': '',

    'set_year': '',
    'begin_depth': '',
    'end_depth': '',

    'pipe_class': '',
    'pipe_material': '',
    'pipe_dia': '',

    'direction': '',
    'lenth': '',
    'test_lenth': '',

    'repair_metric': '',
    'maintain_metric': '',
    'tester': '',

    'address': '',
    'date': '',

    'defect_list': [],
    'note':''
}

# fill the 3.3table
def fill_defect(document: docx.Document, pipe_dic: dict = None, save_path=''):
    tmpl = docx.Document(BASE_DIR+r'/template/defect_template.docx')
    doc2 = document
    table = deepcopy(tmpl.tables[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.cell(0, 1).text = pipe_dic['video_name']
    table.cell(0, 4).text = pipe_dic['begin_well']
    table.cell(0, 7).text = pipe_dic['end_well']

    table.cell(1, 1).text = pipe_dic['set_year']
    table.cell(1, 4).text = pipe_dic['begin_depth']
    table.cell(1, 7).text = pipe_dic['end_depth']

    table.cell(2, 1).text = pipe_dic['pipe_class']
    table.cell(2, 4).text = pipe_dic['pipe_material']
    table.cell(2, 7).text = pipe_dic['pipe_dia']

    table.cell(3, 1).text = pipe_dic['direction']
    table.cell(3, 4).text = pipe_dic['lenth']
    table.cell(3, 7).text = pipe_dic['test_lenth']

    table.cell(4, 1).text = pipe_dic['repair_metric']
    table.cell(4, 4).text = pipe_dic['maintain_metric']
    table.cell(4, 7).text = pipe_dic['tester']

    table.cell(5, 1).text = pipe_dic['address']
    table.cell(5, 7).text = pipe_dic['date']

    defect_num = len(pipe_dic['defect_list'])
    print(f'{defect_num} defects')
    defect_row_num = ceil(defect_num / 2)

    if defect_row_num > 1:
        for i in range(1, defect_row_num):
            img_row = deepcopy(table.rows[9]._tr)
            img_name_row = deepcopy(table.rows[10]._tr)
            insertion_row1 = table.rows[9]._tr
            insertion_row1.addnext(img_row)
            insertion_row2 = table.rows[9]._tr
            insertion_row2.addnext(img_name_row)

    if defect_num > 1:
        for i in range(1, defect_num):
            defect_row = deepcopy(table.rows[7]._tr)
            insertion_row = table.rows[7]._tr
            insertion_row.addnext(defect_row)

    paragraph = doc2.paragraphs[-1]
    paragraph._p.addnext(table._tbl)

    table = doc2.tables[-1]
    row_i1 = 7
    row_i2 = 7 + defect_num + 1
    col_i2 = 0

    for i, defect in enumerate(pipe_dic['defect_list']):
        table.cell(row_i1, 0).text = str(defect['distance'])
        table.cell(row_i1, 1).text = defect['name']
        table.cell(row_i1, 2).text = defect['mark']
        table.cell(row_i1, 3).text = defect['grade']
        table.cell(row_i1, 4).text = defect['description']
        table.cell(row_i1, 8).text = f'{i+1}'
        row_i1 += 1

        # img = Image.open(defect['img_path'])
        # height = Cm(table.rows[row_i2].height)
        # width = Cm(table.rows[row_i2].cells[0].width)
        # print(f'{width},{height}')
        # im_resized = img.resize((width, height), resample=Image.BILINEAR)
        # img.show()
        # img.save(r'./report_data/defect_img/'+name)

        run = table.cell(row_i2, col_i2).paragraphs[0].add_run()
        # print(defect['img_path'])
        img_path = os.path.join(save_path, 'defect_images', defect['img_path'])
        print(img_path)
        # picture = run.add_picture(defect['img_path'])
        picture = run.add_picture(img_path)
        picture.width = Cm(8.47)
        picture.height = Cm(6.35)

        table.cell(row_i2 + 1, col_i2).text = f'图片{i+1}'

        if col_i2 == 0:
            col_i2 = 5
        elif col_i2 == 5:
            col_i2 = 0
            row_i2 += 2

    # paragraph.add_run()
    # paragraph.runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    document.add_page_break()


if __name__ == "__main__":
    path = r'D:\code\cctv_demo\template\Template1.docx'
    # generate_docx(path)
    # doc = docx.Document(path)

    # # 打印文档中所有段落的文本内容
    # for para in doc.paragraphs:
    #     print(para.text)

    # 打印文档中所有表格的内容
    # for j, table in enumerate(doc.tables):
    #     print(f'table {j}')
    #     for i, row in enumerate(table.rows):
    #         row_text = [cell.text for cell in row.cells]
    #         print("|".join(row_text))

    # table3 = []
    # for i,row in enumerate(doc.tables[19]):
    #     table3.append(row)



    dic = {
        'video_name': '1sf',
        'begin_well': 'df',
        'end_well': 'fds',

        'set_year': 'sdf',
        'begin_depth': 'df',
        'end_depth': 'dfs',

        'pipe_class': 'sdfa',
        'pipe_material': 'asdf',
        'pipe_dia': 'sdf',

        'direction': 'sdf',
        'lenth': '',
        'test_lenth': 'sdf',

        'repair_metric': 'asdf',
        'maintain_metric': 'sdf',
        'tester': 'asfd',

        'address': 'sdf',
        'date': 'sdf',

        'defect_list': [],
        'note': ''
    }

    defect_dic1 = {
        'distance': 'sdf',
        'name': 'sdf',
        'mark': 'f',
        'grade': 's',
        'description': 'sdf',
        'img_path': r'./report_data/defect_img/00000088.png',

    }

    defect_dic2 = {
        'distance': 'ssf',
        'name': '22',
        'mark': '33',
        'grade': '44',
        'description': 's5f',
        'img_path': r'./report_data/defect_img/00000092.png',

    }

    defect_dic3 = {
        'distance': 'sasgf',
        'name': '2dfg',
        'mark': '3fgh',
        'grade': '4jg',
        'description': 's5ghj',
        'img_path': r'./report_data/defect_img/00000096.png',

    }

    defect_list = []
    defect_list.append(defect_dic1)
    defect_list.append(defect_dic2)
    defect_list.append(defect_dic3)

    dic['defect_list'] = defect_list


    doc = docx.Document(path)
    fill_defect(document=doc, pipe_dic=dic)

    doc.save('./a.docx')
